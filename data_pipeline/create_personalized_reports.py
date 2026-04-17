import logging
import os
import pandas as pd

import inflect # used to convert numbers to words

from dataclasses import dataclass
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from .tools.consts import module_names, get_microskill_description
from .create_survey_variables import get_variable_name

def hex_to_rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

@dataclass(frozen=True)
class TextStyling:
    font_size: float
    color: str
    font_family: str = "Calibri"
    bold: bool = False
    italic: bool = False

def add_styled_paragraph(doc: Document, text: str, style: TextStyling):
    p = doc.add_paragraph()

    run = p.add_run(text)
    run.font.name = style.font_family
    run.font.size = Pt(style.font_size)
    run.bold = style.bold
    run.italic = style.italic
    
    if style.color:
        run.font.color.rgb = hex_to_rgb(style.color)

    #print(type(p)) # <class 'docx.text.paragraph.Paragraph'>
    return p

def create_personalized_reports(final_data_fp: str, microskill_fp: str, output_dir: str, logger: logging.Logger):
    """
    For each student in final data file, create a document covering all progress for all modules.
    """

    report_dir = os.path.join(output_dir, "personalized_reports")
    os.makedirs(report_dir, exist_ok=True)


    # Maps columns (variable) to meanings
    microskill = pd.read_csv(microskill_fp, usecols=["Variable", "Module", "Phase", "Full Question Text", "Module Number", "Microskill Number"])
    modules: list[int] = sorted(microskill["Module Number"].dropna().drop_duplicates().astype(int).to_list())
    microskills: list[int] = sorted(microskill["Microskill Number"].dropna().drop_duplicates().astype(int).to_list())

    overall_max_module = pd.to_numeric(microskill["Module Number"], errors="coerce").max()

    # This dataframe is one row per person
    final_data = pd.read_csv(final_data_fp)

    logger.info(f" ----- CREATING ALL PERSONALIZED REPORTS FOR {final_data["name"].nunique():,} STUDENTS -----")

    for idx, row in final_data.iterrows():
        student_name: str = row["name"]
        document_fp = os.path.join(report_dir, f"{student_name.replace(" ", "_")}_Personalized_Report.docx")

        def has_value(x):
            return not (pd.isna(x) or x == 0)

        max_module = 0
        for module_num in modules:
            pre = row[get_variable_name("mean", module_num, "Pre")]
            post = row[get_variable_name("mean", module_num, "Post")]

            if has_value(pre) and has_value(post):
                max_module = max(max_module, module_num)
    
        # This engine can convert numbers to words
        inflect_engine = inflect.engine()
        max_module_str = inflect_engine.number_to_words(max_module)

        first_name = student_name.split(" ")[0]

        doc = Document()

        logger.info(f"Creating report for {student_name}")

        # 1. There is a tiny piece of text at the top
        first_text = "AI PASSPORT — PERSONALIZED LEARNING REPORT"
        first_text_style = TextStyling(font_size=9, color="#515151", bold=True)
        add_styled_paragraph(doc=doc, text=first_text, style=first_text_style)

        # 2. The person's name as a header
        name_text = student_name
        name_text_style = TextStyling(font_size=18, color="#1F3A5F", bold=True)
        add_styled_paragraph(doc=doc, text=name_text, style=name_text_style)

        # 3. An italic message showing the module number completed
        if max_module <= 0:
            module_completer_text = "Limited engagement"
        else:
            module_completer_text = f"Module {max_module} completer"
            # Another example: Engaged through Module 5 (didnt really do much in module 5, but did all the rest.)
    
        module_summary_text = f"Spring 2026 Cohort  • {module_completer_text}"
        module_sumamry_style = TextStyling(font_size=10, color="#555555", italic=True)
        add_styled_paragraph(doc=doc, text=module_summary_text, style=module_sumamry_style)

        # 4. A small paragraph sumarizing what the report is
        main_text = f"Dear {first_name}, this report reflects the full arc of your work through the Spring 2026 AI Passport program. You completed all {max_module_str} modules end-to-end. Below is what your own pre/post data tells us about how your confidence on the seven microskills of each module shifted across the semester."
        main_style = TextStyling(font_size=10.5, color="#000000")
        add_styled_paragraph(doc=doc, text=main_text, style=main_style)
        
        # 5. For each module in order, create a table with 4 rows and 8 columns. There are 7 microskills per module.
            # The table has columns: 'Microskill', 'Before', 'After', and 'Change'
                # Microskill: a description of the microskill for that module
                # Before: The 'pre' survey type's answer for the microskill. 'M1_Pre_MS1' for module 1 microskill 1
                # After: The 'post' survey type's answer for the microskill. 'M1_Post_MS1' for module 1 microskill 1
                # Change: The delta value (integer). 'M1_Delta_MS1' for module 1 microskill 1. 
                    # If positive, add a '+' before the number and make the cell '#D0E9D8' color
                    # If negative, keep the '-' and make the cell '#DBE2D8' color.
                # If the module pre survey is null, do not create the table
                # If the post survey is null, create the table but have the values for 'After' and 'Change' as '--'
        for module_num in modules:
            pre_survey_mean_col = get_variable_name(variable_type="mean", module_num=module_num, survey_type="Pre")
            pre_survey_mean = row[pre_survey_mean_col]

            post_survey_mean_col = get_variable_name(variable_type="mean", module_num=module_num, survey_type="Post")
            post_survey_mean = row[post_survey_mean_col]

            module_delta_mean_col = get_variable_name(variable_type="delta_mean", module_num=module_num)
            module_delta_mean = row[module_delta_mean_col]

            # Module tables are only created if the pre survey is taken, at the minimum
            if not has_value(pre_survey_mean) and not has_value(post_survey_mean):
                continue
            
            pre_survey_mean = round(pre_survey_mean, 1)
            post_survey_mean = round(post_survey_mean, 1)
            module_delta_mean = round(module_delta_mean, 1)

            # 5.1. The table header
            table_header_text = module_names.get(module_num)
            table_header_style = TextStyling(font_size=12, color="#1F3A5F", bold=True)
            add_styled_paragraph(doc=doc, text=table_header_text, style=table_header_style)

            # 5.2. The table
            rows = len(microskills) + 1
            table_cols = ["Microskill", "Before", "After", "Change"]

            # 5.2a. Create the table and the headers
            table = doc.add_table(rows=rows, cols=len(table_cols))
            table.style = "Table Grid"
            for j, table_col in enumerate(table_cols):
                cell = table.rows[0].cells[j]
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(table_col)
                run.bold = True
                run.font.size = Pt(10.5)
            
            # 5.2b. Color in the odd rows with a light blue color
            for i in range(1, len(table.rows)):
                if i % 2 == 1:
                    for cell in table.rows[i].cells:
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{'#D9E1F2'.lstrip("#")}"/>')
                        cell._element.get_or_add_tcPr().append(shading)

            # 5.2c. Fill in microskill text
            for i, microskill_num in enumerate(microskills):
                microskill_text = get_microskill_description(module_num=module_num, microskill_num=microskill_num)

                before_value_col = get_variable_name(variable_type="microskill", module_num=module_num, microskill_num=microskill_num, survey_type="Pre")
                after_value_col = get_variable_name(variable_type="microskill", module_num=module_num, microskill_num=microskill_num, survey_type="Post")
                delta_value_col = get_variable_name(variable_type="delta_microskill", module_num=module_num, microskill_num=microskill_num)

                before_value: str = row[before_value_col]
                after_value: str | None = row[after_value_col]
                delta_value: int | None = row[delta_value_col]

                if pd.isna(after_value):
                    after_value = "—"

                if pd.isna(delta_value):
                    delta_value = "—"
                else:
                    delta_value = int(delta_value)

                row_cells = table.rows[microskill_num].cells
                row_cells[0].text = microskill_text
                # First column needs to be bold
                for run in row_cells[0].paragraphs[0].runs:
                    run.bold = True

                row_cells[1].text = str(before_value)
                row_cells[2].text = str(after_value)

                if isinstance(delta_value, str):
                    row_cells[3].text = str(delta_value)
                
                elif isinstance(delta_value, int):
                    sign = "+" if delta_value > 0 else ""
                    row_cells[3].text = f"{sign}{delta_value}"
                    
                    if delta_value > 0:
                        color = "#D0E9D8"
                    elif delta_value < 0:
                        color = "#FFD8D6"
                    else:
                        color = None

                    if color:
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color.lstrip("#")}"/>')
                        row_cells[3]._element.get_or_add_tcPr().append(shading)
                
                else:
                    logger.warning(f"Delta value ({delta_value_col}) is not str or int, but instead {type(delta_value)}: {delta_value}")
        
            # 5.2d. Fix the table font not being 10.5, but instead being 11
            for table_row in table.rows:
                for cell in table_row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Calibri"
                            run.font.size = Pt(10.5)

            # 5.3 Add a small label displaying means and delta
            if not isinstance(delta_value, str):
                label_text = f"Average across microskills: {pre_survey_mean} → {post_survey_mean}   (Δ {"+" if module_delta_mean > 0 else ""}{module_delta_mean} on the 1-5 confidence scale)"
                label_style = TextStyling(font_size=9, color="#555555", italic=True)
                add_styled_paragraph(doc=doc, text=label_text, style=label_style)
            else:
                doc.add_paragraph()

        doc.add_paragraph()

        # 6. Last header
        last_header_text = "What this means, and what comes next"
        last_header_style = TextStyling(font_size=11, color="#1F3A5F", bold=True)
        add_styled_paragraph(doc=doc, text=last_header_text, style=last_header_style)

        # 7. Last Paragraph
        if max_module < overall_max_module:
            last_paragraph_text = (
                "You are invited to continue on our Rolling Completion Track: 90 days of Canvas access to the remaining modules, "
                "no required live sessions, and the same completion certificate as cohort completers. For each module you finish, "
                "you may also schedule an optional 20-minute conversation with that module's faculty lead. If timing or life made "
                "this spring difficult, this is a low-pressure way to finish on your own terms. Reply to the accompanying email and "
                "we will extend your access."
            )
        else:
            last_paragraph_text = (
                "The pattern in your pre/post data mirrors what we are seeing across Spring 2026 completers: meaningful, durable gains "
                "on the microskills the program was designed to teach. You are one of a small group that carried the full arc through. "
                "We would like to hear from you about what worked and what did not, and — if useful — to schedule a 20-minute conversation "
                "with any module faculty member about applying these skills to a specific project of yours."
            )
        
        add_styled_paragraph(doc=doc, text=last_paragraph_text, style=main_style)

        doc.add_paragraph()

        # 8. Last Message
        # last_message_text = "Brittney, thank you for being committed to the program and your learning about AI! I hope you can apply skill learned to your next project and use the community you got to know though this program as future collaborators."
        # last_message_style = TextStyling(font_size=9, color="#888888", italic=True)
        # add_styled_paragraph(doc=doc, text=last_message_text, style=last_message_style)
        # add_styled_paragraph(doc=doc, text="Azra Bihorac", style=last_message_style)
        doc.save(document_fp)
